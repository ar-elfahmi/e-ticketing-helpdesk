"""Inject HTTP endpoint specifications into the .docx documentation."""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

import sys, os
PATH_IN = 'Dokumentasi_API_E_Ticketing_Helpdesk.docx'
PATH_OUT = sys.argv[1] if len(sys.argv) > 1 else PATH_IN
doc = Document(PATH_IN)
body_parent = doc.paragraphs[0]._parent

# ---------- helpers ----------

def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def _shade_row(row, color='D9E2F3'):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

def insert_p(prev_elem, text='', bold=False, mono=False, style='Normal', indent=False):
    new_p = OxmlElement('w:p')
    prev_elem.addnext(new_p)
    para = Paragraph(new_p, body_parent)
    try:
        para.style = doc.styles[style]
    except KeyError:
        pass
    if indent:
        para.paragraph_format.left_indent = Cm(0.6)
    if text:
        run = para.add_run(text)
        if bold:
            run.bold = True
        if mono:
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
    return new_p

def insert_label_value(prev_elem, label, value, mono_value=False):
    new_p = OxmlElement('w:p')
    prev_elem.addnext(new_p)
    para = Paragraph(new_p, body_parent)
    para.style = doc.styles['Normal']
    r1 = para.add_run(label)
    r1.bold = True
    if value:
        r2 = para.add_run(value)
        if mono_value:
            r2.font.name = 'Consolas'
            r2.font.size = Pt(9.5)
    return new_p

def insert_json_block(prev_elem, json_text):
    new_p = OxmlElement('w:p')
    prev_elem.addnext(new_p)
    para = Paragraph(new_p, body_parent)
    try:
        para.style = doc.styles['Normal']
    except KeyError:
        pass
    para.paragraph_format.left_indent = Cm(0.6)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(4)
    lines = json_text.split('\n')
    for i, line in enumerate(lines):
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        if i < len(lines) - 1:
            run.add_break()
    return new_p

def insert_table(prev_elem, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    # body
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.add_run(str(val))
    _set_table_borders(table)
    _shade_row(table.rows[0])
    prev_elem.addnext(table._tbl)
    return table._tbl

# Capture anchors BEFORE any insertion
paras = doc.paragraphs
A = {
    'BAB_V_INTRO': paras[66],
    '5.1.1': paras[70],
    '5.1.2': paras[73],
    '5.1.3': paras[75],
    '5.1.4': paras[77],
    '5.2.1': paras[81],
    '5.2.2': paras[83],
    '5.2.3': paras[86],
    '5.2.4': paras[88],
    '5.2.5': paras[90],
    '5.2.6': paras[92],
    '5.2.7': paras[94],
    '5.3.1': paras[97],
    '5.3.2': paras[99],
    '5.4.1': paras[102],
    '5.4.2': paras[104],
    '5.5.1': paras[107],
    '5.5.2': paras[109],
}

# Sanity check anchors
for key, p in A.items():
    print(f'{key}: "{p.text[:60]}..."')

# ---------- per-endpoint helper ----------

SPEC_LABEL = 'Spesifikasi Teknis'

def header_block(anchor_para, method, path, ref):
    a = anchor_para._element
    a = insert_p(a, SPEC_LABEL, bold=True)
    a = insert_label_value(a, 'Method: ', f'{method}')
    a = insert_label_value(a, 'Path: ', path, mono_value=True)
    a = insert_label_value(a, 'Referensi: ', ref, mono_value=True)
    return a

# ---------- BOTTOM-UP INSERTION ----------

# 5.5.2 PUT /users/{userId}
a = header_block(A['5.5.2'], 'PUT', '/users/{userId}', 'ProfileRepository.updateProfile()')
a = insert_label_value(a, 'Path Parameter: ', 'userId (string) — ID pengguna yang akan diperbarui.')
a = insert_label_value(a, 'Body Request: ', '(struktur lengkap UserModel)')
a = insert_json_block(a, '{\n  "id": "u1",\n  "name": "Ari Pratama",\n  "email": "ari.user@example.com",\n  "username": "user",\n  "role": "user",\n  "avatarUrl": null\n}')
a = insert_label_value(a, 'Response Sukses: ', 'Objek UserModel yang sudah diperbarui.')

# 5.5.1 GET /users/{userId}
a = header_block(A['5.5.1'], 'GET', '/users/{userId}', 'ProfileRepository.getUserProfile()')
a = insert_label_value(a, 'Path Parameter: ', 'userId (string) — ID pengguna yang dicari.')
a = insert_label_value(a, 'Response Sukses: ', 'Objek UserModel atau null bila pengguna tidak ditemukan.')

# 5.4.2 PATCH /notifications/{notificationId}/read
a = header_block(A['5.4.2'], 'PATCH', '/notifications/{notificationId}/read', 'NotificationRepository.markAsRead()')
a = insert_label_value(a, 'Path Parameter: ', 'notificationId (string) — ID notifikasi yang ditandai.')
a = insert_label_value(a, 'Body Request: ', 'Tidak ada.')
a = insert_label_value(a, 'Response: ', 'Future<void>. Bila ID tidak ditemukan, operasi diabaikan tanpa pesan kesalahan.')

# 5.4.1 GET /notifications
a = header_block(A['5.4.1'], 'GET', '/notifications', 'NotificationRepository.getNotifications()')
a = insert_label_value(a, 'Query Parameter: ', 'userId (string, opsional) — filter notifikasi untuk pengguna tertentu.')
a = insert_label_value(a, 'Response Sukses: ', 'Array NotificationModel, diurutkan dari yang terbaru (createdAt desc).')
a = insert_json_block(a, '[\n  {\n    "id": "n1",\n    "title": "Status tiket diperbarui",\n    "body": "Tiket TK-002 berubah menjadi In Progress.",\n    "type": "ticket",\n    "isRead": false,\n    "createdAt": "2026-06-04T09:15:00.000Z",\n    "ticketId": "t2",\n    "userId": "u1"\n  }\n]')

# 5.3.2 GET /dashboard/recent-tickets
a = header_block(A['5.3.2'], 'GET', '/dashboard/recent-tickets', 'DashboardRepository.getRecentTickets()')
a = insert_label_value(a, 'Query Parameter:', '')
a = insert_table(a,
    ['Nama', 'Tipe', 'Default', 'Keterangan'],
    [
        ['reporterId', 'string', 'null', 'Filter tiket milik pelapor tertentu.'],
        ['limit', 'integer', '5', 'Jumlah maksimum tiket yang dikembalikan.'],
    ])
a = insert_label_value(a, 'Response Sukses: ', 'Array TicketModel, diurutkan berdasarkan updatedAt (desc).')

# 5.3.1 GET /dashboard/stats
a = header_block(A['5.3.1'], 'GET', '/dashboard/stats', 'DashboardRepository.getStats()')
a = insert_label_value(a, 'Query Parameter: ', 'reporterId (string, opsional) — bila diisi, statistik dibatasi pada tiket milik reporter tersebut.')
a = insert_label_value(a, 'Response Sukses:', '')
a = insert_json_block(a, '{\n  "totalTickets": 12,\n  "openCount": 4,\n  "inProgressCount": 4,\n  "closedCount": 4\n}')

# 5.2.7 POST /tickets/{ticketId}/comments
a = header_block(A['5.2.7'], 'POST', '/tickets/{ticketId}/comments', 'TicketRepository.addComment()')
a = insert_label_value(a, 'Path Parameter: ', 'ticketId (string) — ID tiket tempat komentar ditambahkan.')
a = insert_label_value(a, 'Body Request:', '')
a = insert_json_block(a, '{\n  "userId": "u1",\n  "userName": "Ari Pratama",\n  "content": "Update terbaru dari pelapor."\n}')
a = insert_label_value(a, 'Response Sukses: ', 'Objek CommentModel baru. Atribut id dan createdAt di-generate oleh sistem.')

# 5.2.6 GET /tickets/{ticketId}/comments
a = header_block(A['5.2.6'], 'GET', '/tickets/{ticketId}/comments', 'TicketRepository.getComments()')
a = insert_label_value(a, 'Path Parameter: ', 'ticketId (string) — ID tiket yang dilihat komentarnya.')
a = insert_label_value(a, 'Response Sukses: ', 'Array CommentModel, diurutkan kronologis dari yang paling lama (createdAt asc).')
a = insert_json_block(a, '[\n  {\n    "id": "c1",\n    "ticketId": "t1",\n    "userId": "u1",\n    "userName": "Ari Pratama",\n    "content": "Mohon dibantu follow up untuk tiket ini.",\n    "createdAt": "2026-06-03T09:00:00.000Z"\n  }\n]')

# 5.2.5 PATCH /tickets/{ticketId}/assign
a = header_block(A['5.2.5'], 'PATCH', '/tickets/{ticketId}/assign', 'TicketRepository.assignTicket()')
a = insert_label_value(a, 'Path Parameter: ', 'ticketId (string) — ID tiket yang akan ditugaskan.')
a = insert_label_value(a, 'Body Request:', '')
a = insert_json_block(a, '{\n  "assigneeId": "h1",\n  "updatedBy": "Raka Admin"\n}')
a = insert_label_value(a, 'Response Sukses: ', 'Objek TicketModel dengan assigneeId terisi dan entri history baru.')
a = insert_label_value(a, 'Response Gagal: ', 'null bila tiket tidak ditemukan.')

# 5.2.4 PATCH /tickets/{ticketId}/status
a = header_block(A['5.2.4'], 'PATCH', '/tickets/{ticketId}/status', 'TicketRepository.updateStatus()')
a = insert_label_value(a, 'Path Parameter: ', 'ticketId (string) — ID tiket yang akan diperbarui.')
a = insert_label_value(a, 'Body Request:', '')
a = insert_json_block(a, '{\n  "status": "inProgress",\n  "updatedBy": "Nadia Helpdesk",\n  "note": "Sedang ditangani"\n}')
a = insert_label_value(a, 'Catatan: ', 'Field status dan updatedBy wajib. Field note bersifat opsional.')
a = insert_label_value(a, 'Response Sukses: ', 'Objek TicketModel terbaru dengan history yang sudah ditambahkan.')
a = insert_label_value(a, 'Response Gagal: ', 'null bila tiket tidak ditemukan.')

# 5.2.3 POST /tickets
a = header_block(A['5.2.3'], 'POST', '/tickets', 'TicketRepository.createTicket()')
a = insert_label_value(a, 'Body Request:', '')
a = insert_json_block(a, '{\n  "title": "Judul singkat masalah",\n  "description": "Deskripsi detail masalah",\n  "category": "Hardware",\n  "priority": "medium",\n  "reporterId": "u1",\n  "attachments": []\n}')
a = insert_label_value(a, 'Catatan: ', 'Field title, description, category, priority, reporterId wajib. Field attachments opsional dengan default array kosong.')
a = insert_label_value(a, 'Response Sukses:', '')
a = insert_json_block(a, '{\n  "id": "t16",\n  "ticketNumber": "TK-016",\n  "title": "Judul singkat masalah",\n  "description": "Deskripsi detail masalah",\n  "category": "Hardware",\n  "priority": "medium",\n  "status": "open",\n  "createdAt": "2026-06-04T10:00:00.000Z",\n  "updatedAt": "2026-06-04T10:00:00.000Z",\n  "reporterId": "u1",\n  "assigneeId": null,\n  "attachments": [],\n  "history": [\n    {\n      "status": "open",\n      "changedBy": "Pelapor",\n      "changedAt": "2026-06-04T10:00:00.000Z",\n      "note": "Tiket dibuat"\n    }\n  ]\n}')

# 5.2.2 GET /tickets/{ticketId}
a = header_block(A['5.2.2'], 'GET', '/tickets/{ticketId}', 'TicketRepository.getTicketById()')
a = insert_label_value(a, 'Path Parameter: ', 'ticketId (string) — ID tiket internal, contoh t1.')
a = insert_label_value(a, 'Response Sukses: ', 'Objek TicketModel lengkap (lihat skema pada endpoint /tickets).')
a = insert_label_value(a, 'Response Gagal: ', 'null bila tiket tidak ditemukan.')

# 5.2.1 GET /tickets
a = header_block(A['5.2.1'], 'GET', '/tickets', 'TicketRepository.getTickets()')
a = insert_label_value(a, 'Query Parameter:', '')
a = insert_table(a,
    ['Nama', 'Tipe', 'Default', 'Keterangan'],
    [
        ['status', 'string (open | inProgress | closed)', 'null', 'Filter berdasarkan status tiket.'],
        ['query', 'string', '""', 'Pencarian berdasarkan ticketNumber, title, atau category (case-insensitive).'],
        ['page', 'integer', '1', 'Halaman paginasi.'],
        ['limit', 'integer', '10', 'Jumlah item per halaman.'],
        ['reporterId', 'string', 'null', 'Filter tiket milik pelapor tertentu.'],
    ])
a = insert_label_value(a, 'Response Sukses: ', 'Array TicketModel, diurutkan berdasarkan createdAt (desc).')
a = insert_json_block(a, '[\n  {\n    "id": "t1",\n    "ticketNumber": "TK-001",\n    "title": "Printer lantai 2 tidak bisa cetak",\n    "description": "Printer menampilkan kertas macet padahal tray kosong.",\n    "category": "Hardware",\n    "priority": "medium",\n    "status": "open",\n    "createdAt": "2026-06-03T08:00:00.000Z",\n    "updatedAt": "2026-06-03T14:00:00.000Z",\n    "reporterId": "u1",\n    "assigneeId": "h1",\n    "attachments": [],\n    "history": [\n      {\n        "status": "open",\n        "changedBy": "System",\n        "changedAt": "2026-06-03T08:00:00.000Z",\n        "note": "Tiket dibuat"\n      }\n    ]\n  }\n]')

# 5.1.4 GET /auth/me
a = header_block(A['5.1.4'], 'GET', '/auth/me', 'AuthRepository.getCurrentUser()')
a = insert_label_value(a, 'Body Request: ', 'Tidak ada.')
a = insert_label_value(a, 'Response Sukses: ', 'Objek UserModel (lihat skema pada endpoint /auth/login) atau null bila belum ada sesi aktif.')

# 5.1.3 POST /auth/logout
a = header_block(A['5.1.3'], 'POST', '/auth/logout', 'AuthRepository.logout()')
a = insert_label_value(a, 'Body Request: ', 'Tidak ada.')
a = insert_label_value(a, 'Response: ', 'Future<void> (tidak mengembalikan nilai). Implementasi dummy mengatur _currentUser menjadi null.')

# 5.1.2 POST /auth/register
a = header_block(A['5.1.2'], 'POST', '/auth/register', 'AuthRepository.register()')
a = insert_label_value(a, 'Body Request:', '')
a = insert_json_block(a, '{\n  "name": "Nama Lengkap",\n  "email": "user@example.com",\n  "username": "username_baru",\n  "password": "password123"\n}')
a = insert_label_value(a, 'Catatan: ', 'Seluruh field wajib diisi.')
a = insert_label_value(a, 'Response Sukses:', '')
a = insert_json_block(a, '{\n  "id": "u4",\n  "name": "Nama Lengkap",\n  "email": "user@example.com",\n  "username": "username_baru",\n  "role": "user",\n  "avatarUrl": null\n}')
a = insert_label_value(a, 'Response Gagal: ', 'null bila username atau email sudah terpakai. Pesan UI: "Username atau email sudah dipakai."')

# 5.1.1 POST /auth/login
a = header_block(A['5.1.1'], 'POST', '/auth/login', 'AuthRepository.login()')
a = insert_label_value(a, 'Body Request:', '')
a = insert_json_block(a, '{\n  "username": "user",\n  "password": "password123"\n}')
a = insert_label_value(a, 'Response Sukses:', '')
a = insert_json_block(a, '{\n  "id": "u1",\n  "name": "Ari Pratama",\n  "email": "ari.user@example.com",\n  "username": "user",\n  "role": "user",\n  "avatarUrl": null\n}')
a = insert_label_value(a, 'Response Gagal: ', 'null bila kredensial tidak valid. Pesan UI: "Username atau password salah."')

# ---------- SUMMARY TABLE (paragraph 66, BAB V intro) ----------
a = A['BAB_V_INTRO']._element
a = insert_p(a, 'Tabel 5.1 Ringkasan Daftar Endpoint', bold=True)
a = insert_table(a,
    ['Sub-bab', 'Method', 'Path', 'Operasi Repository'],
    [
        ['5.1.1', 'POST',  '/auth/login',                              'AuthRepository.login()'],
        ['5.1.2', 'POST',  '/auth/register',                           'AuthRepository.register()'],
        ['5.1.3', 'POST',  '/auth/logout',                             'AuthRepository.logout()'],
        ['5.1.4', 'GET',   '/auth/me',                                 'AuthRepository.getCurrentUser()'],
        ['5.2.1', 'GET',   '/tickets',                                 'TicketRepository.getTickets()'],
        ['5.2.2', 'GET',   '/tickets/{ticketId}',                      'TicketRepository.getTicketById()'],
        ['5.2.3', 'POST',  '/tickets',                                 'TicketRepository.createTicket()'],
        ['5.2.4', 'PATCH', '/tickets/{ticketId}/status',               'TicketRepository.updateStatus()'],
        ['5.2.5', 'PATCH', '/tickets/{ticketId}/assign',               'TicketRepository.assignTicket()'],
        ['5.2.6', 'GET',   '/tickets/{ticketId}/comments',             'TicketRepository.getComments()'],
        ['5.2.7', 'POST',  '/tickets/{ticketId}/comments',             'TicketRepository.addComment()'],
        ['5.3.1', 'GET',   '/dashboard/stats',                         'DashboardRepository.getStats()'],
        ['5.3.2', 'GET',   '/dashboard/recent-tickets',                'DashboardRepository.getRecentTickets()'],
        ['5.4.1', 'GET',   '/notifications',                           'NotificationRepository.getNotifications()'],
        ['5.4.2', 'PATCH', '/notifications/{notificationId}/read',     'NotificationRepository.markAsRead()'],
        ['5.5.1', 'GET',   '/users/{userId}',                          'ProfileRepository.getUserProfile()'],
        ['5.5.2', 'PUT',   '/users/{userId}',                          'ProfileRepository.updateProfile()'],
    ])
# Trailing note paragraph after the table
a = insert_p(a, 'Catatan: Method dan path adalah konvensi RESTful prospektif yang diadopsi sebagai acuan implementasi backend di kemudian hari, mengingat aplikasi saat ini masih berjalan dengan repository dummy in-memory.')

doc.save(PATH_OUT)
print(f'\nDocument saved successfully to: {PATH_OUT}')
